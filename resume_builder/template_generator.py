from docx import Document
from docx.shared import Pt, Inches
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Create a new Word document
doc = Document()

# Define color scheme from the website screenshot provided
deep_blue = RGBColor(
    3, 37, 76
)  # This is an approximation based on the screenshot provided
light_gray = RGBColor(230, 230, 230)

# Set the document title
title = doc.add_heading(level=0)
run = title.add_run("Your Name")
run.font.color.rgb = deep_blue
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add zCore Group logo to the header
section = doc.sections[0]
header = section.header
header_paragraph = header.paragraphs[0]
run = header_paragraph.add_run()
# Adjust the path for the logo.png to a known existing location or ensure it exists
logo_path = "resume_builder/logo.png"
run.add_picture(logo_path, width=Inches(2))

# Add contact information placeholder
contact_info = doc.add_paragraph()
contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = contact_info.add_run(
    "Your Role | Email: youremail@example.com | Phone: (123) 456-7890 | LinkedIn: linkedin.com/in/yourprofile"
)
run.font.color.rgb = deep_blue

# Add a line break
doc.add_paragraph()

# Add a placeholder for the Professional Summary section
heading = doc.add_paragraph(style="Heading 1")
run = heading.add_run("Professional Summary")
run.font.color.rgb = deep_blue

# Now, continue with your Professional Summary content
p = doc.add_paragraph()
summary_run = p.add_run(
    "A brief professional summary that highlights your experience, achievements, and goals. "
    "It should be concise and to the point, providing an overview of your professional background."
)
summary_run.font.color.rgb = light_gray

# Add placeholders for other sections
sections = ["Experience", "Education", "Skills", "Certifications"]

for section in sections:
    heading_paragraph = doc.add_paragraph(style="Heading 1")
    run = heading_paragraph.add_run(section)
    run.font.color.rgb = deep_blue

# Define the file path
file_path = "zCore_Branded_Resume_Template.docx"

# Ensure the directory exists or create it
directory = os.path.dirname(file_path)
if not os.path.exists(directory) and directory != "":
    os.makedirs(directory)

# Save the document
doc.save(file_path)

print(f"Document saved to {file_path}")
